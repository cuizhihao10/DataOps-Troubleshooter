"""把 docs/learning/ 下的章节 Markdown 合成一份 Word 学习文档。

只解析本项目章节实际使用的 Markdown 子集：ATX 标题、围栏代码块、管道表格、无序/有序列表、
引用块与段落，行内支持 **粗体**、`代码` 与链接。刻意不引入 pandoc 或通用 Markdown 库：章节里
大量代码块含中文注释和全角标点，通用转换器的换行与字体回退经常把缩进吃掉，而这里的缩进本身
就是要讲的内容。

输出默认写到 docs/learning/dist/，不能落在 docs/ 根目录——tests/unit/test_documentation_policy.py
用 next(Path("docs").glob("*.docx")) 定位正式设计 DOCX，多一个 docx 会让那条门禁指向错文件。

用法：
    .venv/Scripts/python docs/learning/build_docx.py
    .venv/Scripts/python docs/learning/build_docx.py --output build/学习文档.docx
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOC_TITLE = "DataOps-Troubleshooter 项目学习文档"
DOC_SUBTITLE = "证据驱动的大数据排障 Agent：按开发顺序读完全部源码"
BODY_FONT = "Microsoft YaHei"
CODE_FONT = "Consolas"
CODE_FILL = "F2F2F2"
QUOTE_FILL = "F7F7F7"

# 行内标记：代码放在最前面，否则代码里的 ** 会被当成粗体切开。
_INLINE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))")
_LINK = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")
_HEADING = re.compile(r"^(#{1,4})\s+(.*)$")
_ORDERED = re.compile(r"^(\d+)\.\s+(.*)$")
_TABLE_RULE = re.compile(r"^\|[\s:|-]+\|$")
# 中日韩表意文字，以及全角标点与全角形式：拼接规则对这两类字符不一样。
_IDEOGRAPH = re.compile(r"[㐀-䶿一-鿿]")
_FULLWIDTH = re.compile(r"[　-〿＀-￯]")


def join_wrapped(pieces: list[str]) -> str:
    """把 Markdown 里硬换行的多行拼成一段。

    中文段落在源文件里按 100 列手工折行，一律用空格拼接会出现"第一次 出现的地方"这种夹空格；
    一律不加空格又会把"由 Planner"粘成"由Planner"。规则按本项目自己的排版习惯来：全角标点一侧
    不补空格，两侧都是汉字不补空格，其余（汉字与英文/命令相邻）补一个空格。
    """
    text = ""
    for piece in pieces:
        if not piece:
            # 引用块里允许出现空的 "> " 行，跳过即可，否则下面取首字符会越界。
            continue
        if not text:
            text = piece
            continue
        left, right = text[-1], piece[0]
        if _FULLWIDTH.match(left) or _FULLWIDTH.match(right):
            text += piece
        elif _IDEOGRAPH.match(left) and _IDEOGRAPH.match(right):
            text += piece
        else:
            text += " " + piece
    return text



@dataclass
class Block:
    """解析后的一个 Markdown 块，渲染阶段只认这一种结构，不再回看原文。"""

    kind: str
    level: int = 0
    text: str = ""
    ordered: bool = False
    marker: str = ""
    lines: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


def parse_blocks(source: str) -> list[Block]:
    """把一份章节 Markdown 解析为块列表；围栏代码块内的内容一律按字面量保留。

    先判断围栏状态再判断行首标记，否则代码块里的 Python 注释（`# ...`）会被误认成标题——
    十五份章节里这种行有一百多处，是最容易踩的一个坑。
    """
    blocks: list[Block] = []
    paragraph: list[str] = []
    lines = source.splitlines()
    index = 0

    def flush() -> None:
        if paragraph:
            blocks.append(Block(kind="paragraph", text=join_wrapped(paragraph)))
            paragraph.clear()

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            flush()
            language = stripped[3:].strip()
            index += 1
            body: list[str] = []
            while index < len(lines) and lines[index].strip() != "```":
                body.append(lines[index])
                index += 1
            index += 1
            blocks.append(Block(kind="code", text=language, lines=body))
            continue

        if not stripped or stripped.startswith("<!--"):
            flush()
            index += 1
            continue
        heading = _HEADING.match(stripped)
        if heading:
            flush()
            blocks.append(
                Block(kind="heading", level=len(heading.group(1)), text=heading.group(2).strip())
            )
            index += 1
            continue

        if stripped.startswith("|"):
            flush()
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                row = lines[index].strip()
                if not _TABLE_RULE.match(row):
                    cells = [cell.strip() for cell in row.strip("|").split("|")]
                    rows.append(cells)
                index += 1
            if rows:
                blocks.append(Block(kind="table", rows=rows))
            continue

        if stripped.startswith(">"):
            flush()
            quote: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote.append(lines[index].strip().lstrip(">").strip())
                index += 1
            blocks.append(Block(kind="quote", text=join_wrapped(quote)))
            continue

        ordered = _ORDERED.match(stripped)
        if stripped.startswith(("- ", "* ")) or ordered:
            flush()
            marker = f"{ordered.group(1)}." if ordered else "•"
            body_text = ordered.group(2) if ordered else stripped[2:]
            item = [body_text.strip()]
            index += 1
            # 续行以两个空格缩进，属于同一个列表项；空行或新标记结束该项。
            while index < len(lines):
                nxt = lines[index]
                if not nxt.strip() or not nxt.startswith(("  ", "\t")):
                    break
                if nxt.strip().startswith("```") or nxt.strip().startswith("|"):
                    break
                item.append(nxt.strip())
                index += 1
            blocks.append(
                Block(
                    kind="list",
                    text=join_wrapped(item),
                    ordered=bool(ordered),
                    marker=marker,
                )
            )
            continue
        paragraph.append(stripped)
        index += 1

    flush()
    return blocks


def _set_east_asia(rpr, name: str) -> None:
    """Word 对中英文字体分开取值，只设 w:ascii 会让中文回退成宋体，必须显式写 w:eastAsia。"""
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), name)


def _shade(pr, fill: str) -> None:
    """给段落或单元格加底纹；代码块靠底纹和正文区分，比加边框更耐 Word 版本差异。"""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def configure_styles(document) -> None:
    """统一正文、标题与代码样式。样式一次设好，渲染阶段就不用逐个 run 调字体。"""
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    _set_east_asia(normal.element.get_or_add_rPr(), BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for level, size in ((1, 20), (2, 15), (3, 12.5), (4, 11)):
        style = document.styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("1F3864")
        _set_east_asia(style.element.get_or_add_rPr(), BODY_FONT)
        style.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    code = document.styles.add_style("LearningCode", 1)
    code.base_style = document.styles["Normal"]
    code.font.name = CODE_FONT
    code.font.size = Pt(8.5)
    _set_east_asia(code.element.get_or_add_rPr(), CODE_FONT)
    code.paragraph_format.line_spacing = 1.1
    code.paragraph_format.space_before = Pt(6)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.left_indent = Cm(0.3)
    _shade(code.element.get_or_add_pPr(), CODE_FILL)
def add_inline(paragraph, text: str, *, bold: bool = False) -> None:
    """按行内标记切分并追加 run；`代码` 用等宽字体，链接只保留可读文本。"""
    for piece in _INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            run = paragraph.add_run(piece[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(9)
            _set_east_asia(run.element.get_or_add_rPr(), CODE_FONT)
            run.font.color.rgb = RGBColor.from_string("A31515")
            continue
        if piece.startswith("**") and piece.endswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
            continue
        link = _LINK.match(piece)
        if link:
            # 章节间是相对路径互链，Word 里跳不过去，只保留链接文字加一个来源提示。
            run = paragraph.add_run(f"{link.group(1)}（{link.group(2)}）")
            run.bold = bold
            continue
        run = paragraph.add_run(piece)
        run.bold = bold


def render_code(document, block: Block) -> None:
    """一个代码块渲染成一个段落 + 软换行，这样底纹是连续的一整片，不会被段间距切断。"""
    paragraph = document.add_paragraph(style="LearningCode")
    paragraph.paragraph_format.keep_together = True
    if block.text:
        label = paragraph.add_run(f"[{block.text}]")
        label.font.color.rgb = RGBColor.from_string("808080")
        paragraph.add_run().add_break(WD_BREAK.LINE)
    for offset, line in enumerate(block.lines):
        if offset:
            paragraph.runs[-1].add_break(WD_BREAK.LINE)
        paragraph.add_run(line or " ")


def render_table(document, block: Block) -> None:
    """表格用 Table Grid（python-docx 默认模板自带），首行加粗并加底纹。"""
    width = max(len(row) for row in block.rows)
    table = document.add_table(rows=len(block.rows), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(block.rows):
        for column in range(width):
            cell = table.cell(row_index, column)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            value = row[column] if column < len(row) else ""
            add_inline(paragraph, value, bold=row_index == 0)
            if row_index == 0:
                _shade(cell._tc.get_or_add_tcPr(), "EDEDED")
    document.add_paragraph()
def render_block(document, block: Block, *, first_chapter: bool) -> None:
    """把一个块写进文档。章标题（level 1）之前分页，除了整本的第一章。"""
    if block.kind == "heading":
        if block.level == 1 and not first_chapter:
            document.add_page_break()
        paragraph = document.add_paragraph(style=f"Heading {block.level}")
        add_inline(paragraph, block.text)
        return
    if block.kind == "code":
        render_code(document, block)
        return
    if block.kind == "table":
        render_table(document, block)
        return
    if block.kind == "quote":
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.6)
        paragraph.paragraph_format.right_indent = Cm(0.3)
        _shade(paragraph.paragraph_format.element.get_or_add_pPr(), QUOTE_FILL)
        add_inline(paragraph, block.text)
        return
    if block.kind == "list":
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.9)
        paragraph.paragraph_format.first_line_indent = Cm(-0.55)
        paragraph.paragraph_format.space_after = Pt(3)
        # 有序列表保留原文里的数字，不用 Word 自动编号：章节里的序号是被正文引用的。
        paragraph.add_run(f"{block.marker} ")
        add_inline(paragraph, block.text)
        return
    add_inline(document.add_paragraph(), block.text)


def add_cover(document, chapter_count: int, revision: str | None) -> None:
    """封面 + 目录域。目录内容由 Word 打开后按 F9 生成，脚本不去伪造页码。"""
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    run = title.add_run(DOC_TITLE)
    run.bold = True
    run.font.size = Pt(26)
    _set_east_asia(run.element.get_or_add_rPr(), BODY_FONT)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(DOC_SUBTITLE)
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor.from_string("595959")

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source = f"由 docs/learning/ 的 {chapter_count} 份章节 Markdown 生成"
    if revision:
        source += f"｜代码版本 {revision}"
    note.add_run(source).font.size = Pt(9)
    document.add_page_break()
    toc_heading = document.add_paragraph(style="Heading 1")
    toc_heading.add_run("目录")
    holder = document.add_paragraph()
    field_element = OxmlElement("w:fldSimple")
    field_element.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    field_run = OxmlElement("w:r")
    field_text = OxmlElement("w:t")
    field_text.text = "（在 Word 中右键此处选择“更新域”生成带页码的目录）"
    field_run.append(field_text)
    field_element.append(field_run)
    holder._p.append(field_element)


def build(source_dir: Path, output: Path, revision: str | None) -> tuple[int, int]:
    """读取章节、渲染并保存，返回（章节数, 块数）供命令行打印。"""
    chapters = sorted(source_dir.glob("*.md"))
    if not chapters:
        raise SystemExit(f"no chapter markdown found under {source_dir}")

    document = Document()
    configure_styles(document)
    document.core_properties.title = DOC_TITLE
    document.core_properties.subject = DOC_SUBTITLE
    add_cover(document, len(chapters), revision)

    total_blocks = 0
    for position, chapter in enumerate(chapters):
        blocks = parse_blocks(chapter.read_text(encoding="utf-8"))
        total_blocks += len(blocks)
        for offset, block in enumerate(blocks):
            render_block(document, block, first_chapter=position == 0 and offset == 0)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return len(chapters), total_blocks


def main() -> None:
    """命令行入口：默认输出到 docs/learning/dist/，避免污染 docs/ 根目录的正式 DOCX。"""
    root = Path(__file__).resolve().parent
    parser = argparse.ArgumentParser(description="build the learning docx from chapter markdown")
    parser.add_argument("--source-dir", type=Path, default=root)
    parser.add_argument(
        "--output",
        type=Path,
        default=root / "dist" / "DataOps-Troubleshooter-项目学习文档.docx",
    )
    parser.add_argument("--code-revision", default=None, help="git sha recorded on the cover page")
    args = parser.parse_args()

    chapters, blocks = build(args.source_dir, args.output, args.code_revision)
    print(f"chapters={chapters} blocks={blocks} output={args.output}")


if __name__ == "__main__":
    main()
